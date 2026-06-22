"""
fetch_empreendimentos.py
Morais Engenharia · Site de Empreendimentos
v3 — Junho 2026

Gera:
  data_emp.json                      — dados para o site
  FRE__<NOME>__RV<N>.xlsx            — Ficha Resumo do Empreendimento
  CARTA_PROPOSTA__<NOME>__RV<N>.docx — Carta Proposta CEF
  MEMORIAL_HABITACAO__<NOME>__RV<N>.docx
  MEMORIAL_INFRAESTRUTURA__<NOME>__RV<N>.docx
  QUADROS_ABNT__<NOME>__RV<N>.xlsx   — Quadros NBR 12.721

Segredos necessários (GitHub Actions Secrets):
  NOTION_TOKEN_EMP  — token do BD EMPREENDIMENTOS
  NOTION_DB_EMP     — ID do BD EMPREENDIMENTOS
  NOTION_TOKEN_UH   — token do BD UNIDADES DO EMPREENDIMENTO
  NOTION_DB_UH      — ID do BD UNIDADES DO EMPREENDIMENTO
"""

import os, json, requests, shutil, re, copy
from datetime import datetime, date

# ── Credenciais ───────────────────────────────────────────────────────────────
NOTION_TOKEN = os.environ.get("NOTION_TOKEN_EMP", "")
NOTION_DB    = os.environ.get("NOTION_DB_EMP", "")

# BD de Unidades — token e DB separados
NOTION_TOKEN_UH = os.environ.get("NOTION_TOKEN_UH", "")
NOTION_DB_UH    = os.environ.get("NOTION_DB_UH", "")

# ── DEBUG TEMPORÁRIO — remover após confirmar token ──────────────────────────
def _debug_token(nome, token):
    if not token:
        print(f"  DEBUG {nome}: VAZIO")
    else:
        print(f"  DEBUG {nome}: len={len(token)} inicio='{token[:8]}' fim='{token[-6:]}'")

_debug_token("NOTION_TOKEN_EMP", NOTION_TOKEN)
_debug_token("NOTION_TOKEN_UH",  NOTION_TOKEN_UH)
_debug_token("NOTION_DB_UH",     NOTION_DB_UH)
# ── FIM DEBUG ─────────────────────────────────────────────────────────────────

def _headers(token):
    return {
        "Authorization": f"Bearer {token}",
        "Content-Type":  "application/json",
        "Notion-Version": "2022-06-28",
    }

HEADERS    = _headers(NOTION_TOKEN)
HEADERS_UH = _headers(NOTION_TOKEN_UH)

# ── Helpers de propriedades Notion ───────────────────────────────────────────

def get_texto(prop):
    if not prop: return ""
    t = prop.get("type", "")
    if t == "title":     items = prop.get("title", [])
    elif t == "rich_text": items = prop.get("rich_text", [])
    else: return ""
    return "".join(i.get("plain_text", "") for i in items).strip()

def get_select(prop):
    if not prop: return ""
    s = prop.get("select")
    return s.get("name", "").strip() if s else ""

def get_numero(prop):
    if not prop: return None
    return prop.get("number")

def get_data(prop):
    if not prop: return ""
    d = prop.get("date")
    return d.get("start", "") if d else ""

def get_texto_ou_select(prop):
    """Lê campo que pode ser rich_text OU select (migração Notion)."""
    if not prop: return ""
    t = prop.get("type", "")
    if t == "select":    return get_select(prop)
    if t == "rich_text": return get_texto(prop)
    if t == "title":     return get_texto(prop)
    return ""

def sim_nao(props, *nomes):
    for n in nomes:
        v = get_select(props.get(n) or {})
        if v.upper() in ("SIM", "S", "YES", "TRUE"): return True
        if v.upper() in ("NÃO", "NAO", "N", "NO", "FALSE"): return False
    return False

def get_prop(props, *nomes):
    for n in nomes:
        v = props.get(n)
        if v is not None: return v
    return {}

# ── Busca genérica de páginas ─────────────────────────────────────────────────

def buscar_paginas(db_id, headers, label=""):
    db_id = db_id.replace("ntn_", "", 1) if db_id.startswith("ntn_") else db_id
    print(f"Buscando {label} — DB: {db_id}")

    url = f"https://api.notion.com/v1/databases/{db_id}/query"
    resultados, has_more, cursor = [], True, None

    while has_more:
        body = {"page_size": 100}
        if cursor: body["start_cursor"] = cursor
        resp = requests.post(url, headers=headers, json=body, timeout=30)
        print(f"  HTTP {resp.status_code}")
        if resp.status_code != 200:
            print(f"  ERRO: {resp.text[:500]}"); break
        data = resp.json()
        resultados.extend(data.get("results", []))
        has_more = data.get("has_more", False)
        cursor   = data.get("next_cursor")

    print(f"  → {len(resultados)} registros")

    # Debug: mostra campos da 1ª página
    if resultados:
        print(f"  Campos ({label}):")
        for k, v in resultados[0].get("properties", {}).items():
            t = v.get("type", "?")
            if   t == "select":    val = get_select(v)
            elif t == "title":     val = get_texto(v)
            elif t == "rich_text": val = get_texto(v)
            elif t == "number":    val = str(get_numero(v))
            else:                  val = ""
            print(f"    {t:12} | {k:40} | {val}")

    return resultados

# ── Processar página EMPREENDIMENTOS ─────────────────────────────────────────

def processar_pagina(page):
    props = page.get("properties", {})

    nome = ""
    for v in props.values():
        if v.get("type") == "title":
            nome = get_texto(v); break

    setor       = get_texto_ou_select(get_prop(props, "SETOR", "Setor", "setor"))
    cidade      = get_texto_ou_select(get_prop(props, "CIDADE", "Cidade", "cidade"))
    proponente  = get_texto_ou_select(get_prop(props, "PROPONENTE", "Proponente"))
    doc_prop    = get_texto_ou_select(get_prop(props, "DOC. PROPONENTE", "DOC PROPONENTE", "CNPJ PROPONENTE"))
    construtora = get_texto_ou_select(get_prop(props, "CONSTRUTORA", "Construtora"))
    doc_const   = get_texto_ou_select(get_prop(props, "DOC. CONSTRUTORA", "DOC CONSTRUTORA", "CNPJ CONSTRUTORA"))
    resp_tec    = get_texto_ou_select(get_prop(props, "RESPONSÁVEL TÉCNICO", "RESPONSAVEL TECNICO", "Responsável Técnico"))
    doc_resp    = get_texto_ou_select(get_prop(props, "DOC. RESPONSÁVEL", "DOC RESPONSAVEL", "CPF RESPONSÁVEL"))

    rua         = get_texto(get_prop(props, "RUA", "Rua", "rua", "ENDEREÇO", "Endereço"))
    complemento = get_texto(get_prop(props, "COMPLEMENTO", "Complemento"))
    cep         = get_texto(get_prop(props, "CEP", "Cep", "cep"))
    crea        = get_texto(get_prop(props, "CREA", "CAU/CREA", "CAU"))
    incorporador  = get_texto_ou_select(get_prop(props, "INCORPORADOR", "Incorporador"))
    doc_incorp    = get_texto_ou_select(get_prop(props, "DOC. INCORPORADOR", "DOC INCORPORADOR"))
    tel           = get_texto(get_prop(props, "TEL. CONTATO", "TEL CONTATO", "TELEFONE", "Telefone"))
    email         = get_texto(get_prop(props, "EMAIL", "E-MAIL", "Email"))
    agencia       = get_texto(get_prop(props, "AGÊNCIA", "AGENCIA", "Agência"))

    n_un_v = get_numero(get_prop(props, "Nº DE UNIDADES", "N DE UNIDADES", "N° DE UNIDADES", "UNIDADES", "Nº UNIDADES"))
    n_un   = int(n_un_v) if n_un_v is not None else ""

    prazo_v = get_numero(get_prop(props, "PRAZO PREVISTO", "PRAZO", "Prazo Previsto"))
    prazo   = int(prazo_v) if prazo_v is not None else ""

    area_lote_v  = get_numero(get_prop(props, "ÁREA LOTE", "ÁREA DO LOTE", "AREA LOTE", "AREA DO LOTE", "Área do Lote"))
    area_lote    = area_lote_v if area_lote_v is not None else ""

    area_equiv_v = get_numero(get_prop(props, "ÁREA EQUIVALENTE TOTAL", "ÁREA EQUIVALENTE", "AREA EQUIVALENTE TOTAL", "AREA EQUIVALENTE", "Área Equivalente"))
    area_equiv   = area_equiv_v if area_equiv_v is not None else ""

    valor_terreno_v = get_numero(get_prop(props, "VALOR DO TERRENO", "Valor do Terreno", "VALOR TERRENO"))
    valor_terreno   = valor_terreno_v if valor_terreno_v is not None else ""

    revisao_v = get_numero(get_prop(props, "REVISÃO", "REVISAO", "Revisão", "Revisao"))
    revisao   = int(revisao_v) if revisao_v is not None else 0

    proc  = sim_nao(props, "PROCESSO INICIADO", "PROCESSO INCIADO", "Processo Iniciado")
    apri  = sim_nao(props, "APROVAÇÃO INICIADA", "APROVAÇÃO INCIADA", "APROVACAO INICIADA", "APROVACAO INCIADA", "Aprovação Iniciada")
    aprc  = sim_nao(props, "APROVAÇÃO CONCLUÍDA", "APROVACAO CONCLUIDA", "Aprovação Concluída")
    obra  = sim_nao(props, "OBRA INICIADA", "OBRA INCIADA", "Obra Iniciada")
    obraf = sim_nao(props, "OBRA FINALIZADA", "OBRA FINALZIADA", "Obra Finalizada")

    ini_proc = get_data(get_prop(props, "INÍCIO DO PROCESSO", "INICIO DO PROCESSO", "Início do Processo"))
    ini_apro = get_data(get_prop(props, "INÍCIO DA APROVAÇÃO", "INICIO DA APROVACAO", "Início da Aprovação"))
    ter_apro = get_data(get_prop(props, "TÉRMINO DA APROVAÇÃO", "TERMINO DA APROVACAO"))
    ini_obra = get_data(get_prop(props, "INÍCIO DAS OBRAS", "INICIO DAS OBRAS", "Início das Obras"))
    ter_obra = get_data(get_prop(props, "TÉRMINO DAS OBRAS", "TERMINO DAS OBRAS"))

    if   obraf: sl, sc = "OBRA FINALIZADA",    "finalizada"
    elif obra:  sl, sc = "OBRA INICIADA",       "obra"
    elif aprc:  sl, sc = "APROVAÇÃO CONCLUÍDA", "aprovada"
    elif apri:  sl, sc = "APROVAÇÃO INICIADA",  "aprovando"
    elif proc:  sl, sc = "PROCESSO INICIADO",   "processo"
    else:       sl, sc = "NÃO INICIADO",        "pendente"

    return {
        "id": page["id"], "nome": nome,
        "setor": setor, "cidade": cidade, "rua": rua,
        "complemento": complemento, "cep": cep,
        "status_label": sl, "status_cor": sc,
        "processo_iniciado": proc, "aprovacao_iniciada": apri,
        "aprovacao_concluida": aprc, "obra_iniciada": obra, "obra_finalizada": obraf,
        "inicio_processo": ini_proc, "inicio_aprovacao": ini_apro,
        "termino_aprovacao": ter_apro, "inicio_obras": ini_obra, "termino_obras": ter_obra,
        "proponente": proponente, "doc_proponente": doc_prop,
        "construtora": construtora, "doc_construtora": doc_const,
        "responsavel_tecnico": resp_tec, "crea": crea, "doc_responsavel": doc_resp,
        "incorporador": incorporador, "doc_incorporador": doc_incorp,
        "tel_contato": tel, "email": email, "agencia": agencia,
        "n_unidades": n_un, "prazo_previsto": prazo,
        "area_lote": area_lote, "area_equivalente": area_equiv,
        "valor_terreno": valor_terreno, "revisao": revisao,
    }

# ── Processar página UNIDADES DO EMPREENDIMENTO ───────────────────────────────

def processar_unidade(page):
    """
    Retorna dict com os dados de uma unidade.
    - uh_num: número da UH (nome da página, ex: "1", "2"...)
    - nome_emp: campo NOME DO EMPREENDIMENTO (texto) — chave de ligação
    - area_const_priv, area_const_com: áreas construídas
    - area_desc_priv, area_desc_com: áreas descobertas
    - tipologia: campo TIPOLOGIA (select ou texto)
    """
    props = page.get("properties", {})

    # Nome da página = número da UH
    uh_num = ""
    for v in props.values():
        if v.get("type") == "title":
            uh_num = get_texto(v); break

    nome_emp = get_texto(get_prop(props,
        "NOME DO EMPREENDIMENTO", "NOME DO EMPREENDI...", "nome_empreendimento"))

    tipologia = get_texto_ou_select(get_prop(props, "TIPOLOGIA", "Tipologia"))

    area_const_priv = get_numero(get_prop(props,
        "ÁREA CONSTRUÍDA PRIVATIVA", "AREA CONSTRUIDA PRIVATIVA",
        "Área Construída Privativa")) or 0.0

    area_const_com = get_numero(get_prop(props,
        "ÁREA CONSTRUÍDA COMUM", "AREA CONSTRUIDA COMUM",
        "Área Construída Comum")) or 0.0

    area_desc_priv = get_numero(get_prop(props,
        "ÁREA DESCOBERTA PRIVATIVA", "AREA DESCOBERTA PRIVATIVA",
        "Área Descoberta Privativa")) or 0.0

    area_desc_com = get_numero(get_prop(props,
        "ÁREA DESCOBERTA COMUM", "AREA DESCOBERTA COMUM",
        "Área Descoberta Comum")) or 0.0

    return {
        "uh_num":         uh_num,
        "nome_emp":       nome_emp,
        "tipologia":      tipologia,
        "area_const_priv": area_const_priv,
        "area_const_com":  area_const_com,
        "area_desc_priv":  area_desc_priv,
        "area_desc_com":   area_desc_com,
    }

# ── Helpers de formatação ─────────────────────────────────────────────────────

def nome_arquivo(emp, prefixo, ext):
    """FRE__VALE_DAS_BRISAS__RV0.xlsx"""
    rv   = emp.get("revisao", 0)
    safe = "".join(c for c in emp["nome"] if c.isalnum() or c in " _-")
    safe = safe.strip().replace(" ", "_")
    return f"{prefixo}__{safe}__RV{rv}.{ext}"

def brl_fmt(v):
    if v == "" or v is None: return ""
    return f"R$ {v:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

def data_extenso():
    meses = ["janeiro","fevereiro","março","abril","maio","junho",
             "julho","agosto","setembro","outubro","novembro","dezembro"]
    h = date.today()
    return str(h.day), meses[h.month - 1], str(h.year)

def data_hoje_fmt():
    return date.today().strftime("%d/%m/%Y")

# ── Substituição genérica em Word (.docx) ─────────────────────────────────────

def substituir_docx(doc, mapa):
    """
    Substitui todos os placeholders em parágrafos e tabelas do Word.
    Reconstrói runs do parágrafo inteiro para tratar fragmentação do Word.
    """
    def sub_para(paragraph):
        texto = "".join(r.text for r in paragraph.runs)
        for k, v in mapa.items():
            texto = texto.replace(k, str(v) if v else "")
        if paragraph.runs:
            paragraph.runs[0].text = texto
            for r in paragraph.runs[1:]:
                r.text = ""

    for para in doc.paragraphs:
        sub_para(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    sub_para(para)

# ── Gerar FRE (.xlsx) ─────────────────────────────────────────────────────────

def gerar_fre(emp):
    from openpyxl import load_workbook

    src = None
    for f in sorted(os.listdir("."), reverse=True):
        if f.upper().startswith("FRE") and f.endswith(".xlsx"):
            src = f; break
    if not src:
        print("  AVISO: FRE_v*.xlsx não encontrado"); return None

    hoje = data_hoje_fmt()
    n_un = emp["n_unidades"]
    vt_fmt = brl_fmt(emp.get("valor_terreno", ""))

    mapa_cel = {
        "B7":  emp["nome"],       "B10": emp["rua"],
        "M10": emp["complemento"], "B13": emp["setor"],
        "I13": emp["cidade"],     "N13": emp["cep"],
        "B16": emp["proponente"], "M16": emp["doc_proponente"],
        "B19": emp["construtora"],"M19": emp["doc_construtora"],
        "B22": emp["responsavel_tecnico"], "J22": emp["crea"],
        "M22": emp["doc_responsavel"],
        "B25": emp["incorporador"],"M25": emp["doc_incorporador"],
        "B28": emp["responsavel_tecnico"],"J28": emp["tel_contato"],
        "M28": emp["email"],
        "E36": n_un if n_un != "" else None,
        "N59": emp["prazo_previsto"] if emp["prazo_previsto"] != "" else None,
        "I66": emp["area_lote"] if emp["area_lote"] != "" else None,
        "I67": emp["area_equivalente"] if emp["area_equivalente"] != "" else None,
        "C227": hoje,
    }

    out = nome_arquivo(emp, "FRE", "xlsx")
    shutil.copy2(src, out)
    wb = load_workbook(out)
    ws = wb["FRE"]

    for cel, val in mapa_cel.items():
        if val is None or val == "": continue
        ws[cel].value = val

    # Substitui placeholders de texto em toda a planilha
    for row in ws.iter_rows():
        for cell in row:
            if isinstance(cell.value, str):
                v = cell.value
                v = v.replace("{VALOR DO TERRENO}", vt_fmt)
                v = v.replace("{Nº UNIDADES}", str(n_un) if n_un != "" else "")
                if v != cell.value:
                    cell.value = v

    wb.save(out)
    print(f"  FRE: {out}")
    return out

# ── Gerar Carta Proposta (.docx) ──────────────────────────────────────────────

def remover_protecao_docx(doc):
    """Remove proteção de edição do documento Word (w:documentProtection)."""
    from docx.oxml.ns import qn
    settings = doc.settings.element
    protection = settings.find(qn('w:documentProtection'))
    if protection is not None:
        settings.remove(protection)

def gerar_carta_proposta(emp):
    try:
        from docx import Document
    except ImportError:
        print("  AVISO: python-docx não instalado"); return None

    src = next((f for f in ["CARTA PROPOSTA.docx","CARTA_PROPOSTA.docx"] if os.path.exists(f)), None)
    if not src:
        print(f"  AVISO: {src} não encontrado"); return None

    dia, mes, ano = data_extenso()
    n_un  = str(emp["n_unidades"]) if emp["n_unidades"] != "" else ""
    prazo = str(emp["prazo_previsto"]) if emp["prazo_previsto"] != "" else ""
    vt_fmt = brl_fmt(emp.get("valor_terreno", ""))

    mapa = {
        "{AGÊNCIA}":                 emp.get("agencia", ""),
        "{NOME DO EMPRRENDIMENTO}":  emp["nome"],
        "{NOME DO EMPREENDIMENTO}":  emp["nome"],
        "{CONSTRUTORA}":             emp["construtora"],
        "{DOC. CONSTRUTORA}":        emp["doc_construtora"],
        "{RESPONSÁVEL TÉCNICO}":     emp["responsavel_tecnico"],
        "{TEL. CONTATO}":            emp["tel_contato"],
        "{RUA}":                     emp["rua"],
        "{COMPLEMENTO}":             emp["complemento"],
        "{SETOR}":                   emp["setor"],
        "{CIDADE}":                  emp["cidade"],
        "{CEP}":                     emp["cep"],
        "{Nº DE UNIDADES}":          n_un,
        "{VALOR DO TERRENO}":        vt_fmt,
        # datas fracionadas que o Word cortou
        "HJ":  dia,
        "MES": mes,
        "ANO": ano,
    }

    doc = Document(src)

    # Tratamento especial: {PRAZO} cortado como {PRA pelo Word
    def sub_para_carta(paragraph):
        texto = "".join(r.text for r in paragraph.runs)
        # Normaliza qualquer fragmento de {PRAZO}
        texto = re.sub(r'\{PRA[ZO]*\}?', prazo, texto)
        texto = re.sub(r'\{PRAZO\}', prazo, texto)
        for k, v in mapa.items():
            texto = texto.replace(k, str(v) if v else "")
        if paragraph.runs:
            paragraph.runs[0].text = texto
            for r in paragraph.runs[1:]:
                r.text = ""

    for para in doc.paragraphs:
        sub_para_carta(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    sub_para_carta(para)

    remover_protecao_docx(doc)
    out = nome_arquivo(emp, "CARTA_PROPOSTA", "docx")
    doc.save(out)
    print(f"  Carta Proposta: {out}")
    return out

# ── Gerar Memoriais (.docx) ───────────────────────────────────────────────────

def gerar_memorial(emp, tipo):
    try:
        from docx import Document
    except ImportError:
        print("  AVISO: python-docx não instalado"); return None

    srcs = {
        "HABITACAO":      ["MEMORIAL DESCRITIVO DE HABITACAO.docx","MEMORIAL_DESCRITIVO_DE_HABITACAO.docx","MEMORIAL DESCRITIVO DE HABITACAO.doc","MEMORIAL_DESCRITIVO_DE_HABITACAO.doc"],
        "INFRAESTRUTURA": ["MEMORIAL DESCRITIVO INFRAESTRUTURA.docx","MEMORIAL_DESCRITIVO_INFRAESTRUTURA.docx","MEMORIAL DESCRITIVO INFRAESTRUTURA.doc","MEMORIAL_DESCRITIVO_INFRAESTRUTURA.doc"],
    }
    src = next((n for n in srcs.get(tipo, []) if os.path.exists(n)), None)
    if not src:
        print(f"  AVISO: memorial {tipo} não encontrado"); return None

    dia, mes, ano = data_extenso()
    data_hoje_ext = f"{dia} de {mes} de {ano}"

    mapa = {
        "{PROPORNENTE}":            emp["proponente"],  # typo original
        "{PROPONENTE}":             emp["proponente"],
        "{CONSTRUTORA}":            emp["construtora"],
        "{NOME DO EMPREENDIMENTO}": emp["nome"],
        "{RUA}":                    emp["rua"],
        "{COMPLEMENTO}":            emp["complemento"],
        "{SETOR}":                  emp["setor"],
        "{CIDADE}":                 emp["cidade"],
        "{DATA HOJE}":              data_hoje_ext,
        "{DOC. CONSTRUTORA}":       emp["doc_construtora"],
        "{DOC. PROPONENTE}":        emp["doc_proponente"],
        "{RESPONSÁVEL TÉCNICO}":    emp["responsavel_tecnico"],
        "{CREA}":                   emp["crea"],
    }

    doc = Document(src)
    substituir_docx(doc, mapa)

    prefixo = "MEMORIAL_HABITACAO" if tipo == "HABITACAO" else "MEMORIAL_INFRAESTRUTURA"
    out = nome_arquivo(emp, prefixo, "docx")
    doc.save(out)
    print(f"  Memorial {tipo}: {out}")
    return out

# ── Gerar Quadros ABNT (.xlsx) ────────────────────────────────────────────────

def gerar_quadros_abnt(emp, unidades):
    """
    Gera QUADROS_ABNT__<NOME>__RV<N>.xlsx.
    Correções aplicadas:
    - CAPA: insert_rows antes do TOTAL, fórmulas corretas, merges e formatação preservados
    - QUADRO II: bordas copiadas para todas as linhas de UH
    - QUADRO I: referência CAPA!D{total_row} correta (linha TOTAL dinâmica)
    - QUADROS IV A, IV B, IV B.1: fórmulas corretas por UH, linhas excedentes deletadas
    """
    try:
        from openpyxl import load_workbook
        import copy as _copy
    except ImportError:
        print("  AVISO: openpyxl não instalado"); return None

    src = next((f for f in ["QUADROS ABNT.xlsx","QUADROS_ABNT.xlsx"] if os.path.exists(f)), None)
    if not src:
        print("  AVISO: QUADROS ABNT.xlsx não encontrado"); return None

    uhs = [u for u in unidades if u["nome_emp"].strip().upper() == emp["nome"].strip().upper()]
    def sort_key(u):
        try: return int(u["uh_num"])
        except: return u["uh_num"]
    uhs.sort(key=sort_key)

    n = len(uhs)
    if n == 0:
        print(f"  AVISO: nenhuma UH encontrada para '{emp['nome']}'"); return None

    print(f"  Quadros ABNT: {n} UH(s) para '{emp['nome']}'")

    hoje_fmt = data_hoje_fmt()
    out = nome_arquivo(emp, "QUADROS_ABNT", "xlsx")
    shutil.copy2(src, out)
    wb = load_workbook(out)

    def copiar_estilo(src_cell, dst_cell):
        dst_cell.font          = _copy.copy(src_cell.font)
        dst_cell.fill          = _copy.copy(src_cell.fill)
        dst_cell.border        = _copy.copy(src_cell.border)
        dst_cell.alignment     = _copy.copy(src_cell.alignment)
        dst_cell.number_format = src_cell.number_format

    # ── 1. CAPA ───────────────────────────────────────────────────────────────
    # Template: L9=UH1, L10=TOTAL (merge B10:H10), L11=data (merge B11:H11),
    #           L13=assin, L14=resp (merge B14:H14), L15=crea (merge B15:H15)
    # Para n>1: inserir n-1 linhas na posição 10 (antes do TOTAL)
    ws_capa = wb["CAPA"]

    ws_capa["C3"] = f"{emp['rua']}, {emp['complemento']}; {emp['setor']} - {emp['cidade']}"

    CAPA_TPL = 9    # linha da 1ª UH no template
    CAPA_TOT = 10   # linha TOTAL original

    if n > 1:
        ws_capa.insert_rows(CAPA_TOT, amount=n - 1)
        h = ws_capa.row_dimensions[CAPA_TPL].height
        for i in range(1, n):
            nr = CAPA_TPL + i
            for col in range(1, 10):
                copiar_estilo(ws_capa.cell(row=CAPA_TPL, column=col),
                              ws_capa.cell(row=nr,        column=col))
            ws_capa.row_dimensions[nr].height = h

    total_row = CAPA_TPL + n   # linha TOTAL após inserção

    # Remover merges espúrios que o insert_rows cria nas linhas de UH
    # (openpyxl copia os merges da linha-origem para as novas linhas,
    #  o que bloqueia a escrita de D/E/F/G/H nas UHs extras)
    merges_para_remover = []
    for m in ws_capa.merged_cells.ranges:
        if CAPA_TPL <= m.min_row <= total_row - 1 and str(m) != f"B{total_row}:H{total_row}":
            merges_para_remover.append(str(m))
    for m_str in merges_para_remover:
        ws_capa.unmerge_cells(m_str)

    # Preenche cada UH na CAPA
    for i, uh in enumerate(uhs):
        r          = CAPA_TPL + i
        area_const = round(uh["area_const_priv"] + uh["area_const_com"], 2)
        area_desc  = round(uh["area_desc_priv"]  + uh["area_desc_com"],  2)
        ws_capa.cell(row=r, column=2).value = f"FRAÇÃO  {i+1}"
        ws_capa.cell(row=r, column=3).value = f"CASA {i+1}"
        ws_capa.cell(row=r, column=4).value = area_const
        ws_capa.cell(row=r, column=5).value = area_desc
        ws_capa.cell(row=r, column=6).value = f"=SUM(E{r},D{r})"
        ws_capa.cell(row=r, column=7).value = f"=F{r}"
        ws_capa.cell(row=r, column=8).value = f"=SUM(G{r}/G{total_row})"

    # TOTAL — desmerge se necessário, escreve fórmulas, remerge
    merge_str = f"B{total_row}:H{total_row}"
    # Remover merge existente na linha total (caso já exista do template ou insert)
    ranges_to_remove = [str(m) for m in ws_capa.merged_cells.ranges if str(m) == merge_str]
    for m in list(ws_capa.merged_cells.ranges):
        if str(m) == merge_str:
            ws_capa.unmerge_cells(str(m))
            break

    first, last = CAPA_TPL, CAPA_TPL + n - 1
    ws_capa.cell(row=total_row, column=4).value = f"=SUM(D{first}:D{last})"
    ws_capa.cell(row=total_row, column=5).value = f"=SUM(E{first}:E{last})"
    ws_capa.cell(row=total_row, column=6).value = f"=SUM(F{first}:F{last})"
    ws_capa.cell(row=total_row, column=7).value = f"=SUM(G{first}:G{last})"
    ws_capa.cell(row=total_row, column=8).value = f"=SUM(H{first}:H{last})"

    # Corrigir referências D4 (área terreno = G do total) e D5 (área construída = D do total)
    ws_capa["D4"] = f"=G{total_row}"
    ws_capa["D5"] = f"=D{total_row}"

    # Substituir placeholders no bloco de assinatura (linhas após total_row)
    for r in range(total_row + 1, total_row + 8):
        for col in range(1, 10):
            cell = ws_capa.cell(row=r, column=col)
            if isinstance(cell.value, str):
                v = (cell.value
                     .replace("{CIDADE}", emp.get("cidade", ""))
                     .replace("{DATA HOJE}", hoje_fmt)
                     .replace("{RESPONSÁVEL TÉCNICO}", emp.get("responsavel_tecnico", ""))
                     .replace("{CREA}", emp.get("crea", "")))
                if v != cell.value:
                    cell.value = v

    # ── 2. INFORMAÇÕES PRELIMINARES ───────────────────────────────────────────
    ws_ip = wb["INFORMAÇÕES PRELIMINARES"]
    n_un_str  = str(emp["n_unidades"]) if emp["n_unidades"] != "" else str(n)
    casas_str = f"CASA 01 A CASA {n:02d}" if n > 1 else "CASA 01"

    ws_ip["G10"] = emp.get("responsavel_tecnico", "")
    ws_ip["G11"] = emp.get("crea", "")

    mapa_ip = {
        "{IMCORPORADOR}":           emp.get("incorporador", ""),
        "{DOC. INCORPORADOR}":      emp.get("doc_incorporador", ""),
        "{NOME DO EMPREENDIMENTO}": emp["nome"],
        "{Nº DE UNIDADES}":         n_un_str,
    }
    for row in ws_ip.iter_rows():
        for cell in row:
            if isinstance(cell.value, str) and "{" in cell.value:
                v = cell.value
                for k, val in mapa_ip.items():
                    v = v.replace(k, str(val) if val else "")
                if "CASA 01 A" in v:
                    v = v.replace(f"CASA 01 A {n_un_str}", casas_str)
                if v != cell.value:
                    cell.value = v

    # ── 3. QUADRO I ───────────────────────────────────────────────────────────
    ws_q1 = wb["QUADRO I"]
    ws_q1["C10"] = hoje_fmt
    # C16 = soma da coluna D da CAPA (linha TOTAL dinâmica)
    ws_q1["C16"] = f"=CAPA!D{total_row}"
    # Linha TOTAIS do QI (linha 39 no template) — referenciam CAPA!D{total_row}
    ws_q1["C39"] = f"=CAPA!D{total_row}"
    ws_q1["F39"] = f"=CAPA!D{total_row}"
    ws_q1["G39"] = f"=CAPA!D{total_row}"
    ws_q1["R39"] = f"=CAPA!D{total_row}"
    ws_q1["S39"] = f"=CAPA!D{total_row}"

    # ── 4. QUADRO II ──────────────────────────────────────────────────────────
    # Template: L17=CASA01, L18=linha vazia, L19=TOTAIS, L20=ÁREA REAL GLOBAL
    # Estratégia: deletar L18 vazia, inserir n linhas em L18, copiar bordas para todas
    ws_q2 = wb["QUADRO II"]
    ws_q2["C10"] = hoje_fmt

    Q2_TPL   = 17   # linha CASA01 original
    Q2_VAZIA = 18   # linha vazia — deletar
    Q2_TOT   = 19   # linha TOTAIS original

    ws_q2.delete_rows(Q2_VAZIA)
    # Agora: L17=CASA01, L18=TOTAIS

    ws_q2.insert_rows(Q2_VAZIA, amount=n)
    # Agora: L17=CASA01_tpl, L18..17+n=novas, 18+n=TOTAIS
    # ATENÇÃO: insert_rows(18, n) cria n linhas (18,19,...,17+n),
    # mas escrevemos UHs em 17,18,...,16+n — a linha 17+n fica vazia.
    # Deletar essa linha excedente antes de continuar.
    ws_q2.delete_rows(Q2_TPL + n)
    # Agora: L17..16+n=slots das UHs, L17+n=TOTAIS

    q2_totais      = Q2_TPL + n       # linha TOTAIS (corrigida após deleção)
    q2_area_global = q2_totais + 1    # linha ÁREA REAL GLOBAL

    h2 = ws_q2.row_dimensions[Q2_TPL].height

    for i in range(n):
        r      = Q2_TPL + i
        nr     = r  # alias para clareza
        capa_r = CAPA_TPL + i

        # Copiar bordas/estilo da linha 17 para todas as linhas de UH (inclusive a 1ª)
        for col in range(1, 24):
            copiar_estilo(ws_q2.cell(row=Q2_TPL, column=col),
                          ws_q2.cell(row=nr,      column=col))
        ws_q2.row_dimensions[nr].height = h2

        ws_q2.cell(row=r, column=2).value  = f"CASA {i+1:02d}"
        ws_q2.cell(row=r, column=3).value  = f"=CAPA!D{capa_r}"
        ws_q2.cell(row=r, column=4).value  = f"=CAPA!E{capa_r}"
        ws_q2.cell(row=r, column=6).value  = f"=SUM(C{r}+D{r})"
        ws_q2.cell(row=r, column=7).value  = f"=SUM(C{r}+E{r})"
        ws_q2.cell(row=r, column=11).value = f"=SUM(H{r}+I{r})"
        ws_q2.cell(row=r, column=12).value = f"=SUM(H{r}+J{r})"
        ws_q2.cell(row=r, column=13).value = f"=SUM(G{r}+L{r})"
        ws_q2.cell(row=r, column=14).value = f"=CAPA!H{capa_r}"
        ws_q2.cell(row=r, column=15).value = f"=ROUND('QUADRO II'!N{r}*'QUADRO I'!$M$39,2)"
        ws_q2.cell(row=r, column=16).value = f"=ROUND('QUADRO II'!N{r}*'QUADRO I'!$N$39,2)"
        ws_q2.cell(row=r, column=17).value = f"=ROUND('QUADRO II'!N{r}*'QUADRO I'!$O$39,2)"
        ws_q2.cell(row=r, column=18).value = f"=SUM(O{r}+P{r})"
        ws_q2.cell(row=r, column=19).value = f"=SUM(O{r}+Q{r})"
        ws_q2.cell(row=r, column=20).value = f"=SUM(F{r}+K{r}+R{r})"
        ws_q2.cell(row=r, column=21).value = f"=SUM(M{r}+S{r})"
        ws_q2.cell(row=r, column=22).value = 1

    # Atualiza TOTAIS do Q2
    f2, l2 = Q2_TPL, Q2_TPL + n - 1
    for col_letter in list("CDEFGHIJKLMNOPQRSTU"):
        ws_q2[f"{col_letter}{q2_totais}"] = (
            f"=SUMPRODUCT({col_letter}{f2}:{col_letter}{l2},"
            f"$V{f2}:$V{l2})"
        )
    ws_q2[f"V{q2_totais}"] = f"=SUM(V{f2}:V{l2})"

    # Corrigir referências ÁREA REAL GLOBAL e ÁREA EQUIVALENTE GLOBAL
    ws_q2[f"G{q2_area_global}"] = f"=T{q2_totais}"
    ws_q2[f"S{q2_area_global}"] = f"=U{q2_totais}"

    # ── 5. QUADROS IV A, IV B e IV B.1 ───────────────────────────────────────
    # Template tem linha 14=UH1 e linhas 15+ com #REF! (para até 62 UHs).
    # Estratégia: preencher apenas as linhas 14..13+n com fórmulas corretas,
    # depois deletar as linhas excedentes (14+n até a linha antes de TOTAIS/fim).

    # ── 5a. QUADRO IV A ───────────────────────────────────────────────────────
    ws_4a = wb["QUADRO IV A"]
    # Linha TOTAIS no IV A está na linha 76 (template)
    # Linhas de UH: 14..75 (62 slots)
    IVA_FIRST = 14
    IVA_TOT   = 76   # linha TOTAIS (fixo no template)

    # Preencher fórmulas para as n UHs reais
    for i in range(n):
        row_4a = IVA_FIRST + i
        row_q2 = Q2_TPL + i   # linha correspondente no QUADRO II
        ws_4a.cell(row=row_4a, column=2).value  = f"=('QUADRO II'!B{row_q2})"
        ws_4a.cell(row=row_4a, column=3).value  = f"=('QUADRO II'!U{row_q2})"
        ws_4a.cell(row=row_4a, column=4).value  = f"='QUADRO III'!L32*0.0161290322580645"
        ws_4a.cell(row=row_4a, column=5).value  = f"=('QUADRO II'!N{row_q2})"
        ws_4a.cell(row=row_4a, column=12).value = f"=ROUND(G{row_4a}*$K${IVA_TOT - n + IVA_FIRST},2)"
        ws_4a.cell(row=row_4a, column=13).value = f"=('QUADRO II'!V{row_q2})"
        ws_4a.cell(row=row_4a, column=15).value = f"=SUM(M{row_4a}-N{row_4a})"

    # Deletar linhas excedentes (IVA_FIRST+n até IVA_TOT-1)
    # Calcular quantas linhas excedentes há entre última UH real e TOTAIS
    excesso_4a = IVA_TOT - (IVA_FIRST + n)
    if excesso_4a > 0:
        ws_4a.delete_rows(IVA_FIRST + n, excesso_4a)

    # Linha TOTAIS do IV A após deleção das excedentes
    nova_tot_4a = IVA_FIRST + n   # ex: n=2 → linha 16

    # Corrigir referência $K$ na col L para apontar para nova linha de TOTAIS
    for i in range(n):
        row_4a = IVA_FIRST + i
        ws_4a.cell(row=row_4a, column=12).value = f"=ROUND(G{row_4a}*$K${nova_tot_4a},2)"

    # Atualizar fórmulas da linha TOTAIS do IV A
    ws_4a[f"C{nova_tot_4a}"] = f"=SUM(C{IVA_FIRST}:C{IVA_FIRST+n-1})"
    ws_4a[f"D{nova_tot_4a}"] = f"=SUM(D{IVA_FIRST}:D{IVA_FIRST+n-1})"
    ws_4a[f"E{nova_tot_4a}"] = f"=SUM(E{IVA_FIRST}:E{IVA_FIRST+n-1})"
    ws_4a[f"F{nova_tot_4a}"] = f"=SUMPRODUCT(F{IVA_FIRST}:F{IVA_FIRST+n-1},$O{IVA_FIRST}:$O{IVA_FIRST+n-1})"
    ws_4a[f"M{nova_tot_4a}"] = f"=SUM(M{IVA_FIRST}:M{IVA_FIRST+n-1})"
    ws_4a[f"O{nova_tot_4a}"] = f"=SUM(O{IVA_FIRST}:O{IVA_FIRST+n-1})"

    # ── 5b. QUADRO IV B ───────────────────────────────────────────────────────
    ws_4b = wb["QUADRO IV B"]
    # Template: linha 14=UH1, linhas 15+ com #REF!
    # Não há linha TOTAIS explícita — apenas as UHs seguidas de texto de observação
    IVB_FIRST = 14

    # Encontrar até onde vão as linhas de UH (procurar primeira linha sem border após L14)
    # No template, há linhas até aprox 75 com bordas; calcular o total de slots
    # A linha de texto "Designação da unidade" está acima (L12-L13)
    # Slots disponíveis no template: contar linhas com border a partir de L15
    ivb_slots = 0
    for r in range(IVB_FIRST + 1, 200):
        row_cells = list(ws_4b.iter_rows(min_row=r, max_row=r))[0]
        has_border = any(c.border.left.border_style or c.border.right.border_style
                         or c.border.top.border_style or c.border.bottom.border_style
                         for c in row_cells)
        has_content = any(c.value is not None for c in row_cells)
        if not has_border and not has_content:
            break
        ivb_slots += 1
    total_ivb_slots = 1 + ivb_slots  # inclui linha 14

    # Preencher fórmulas para as n UHs reais
    for i in range(n):
        row_4b = IVB_FIRST + i
        row_q2 = Q2_TPL + i
        capa_r = CAPA_TPL + i
        ws_4b.cell(row=row_4b, column=2).value = f"=('QUADRO II'!B{row_q2})"
        ws_4b.cell(row=row_4b, column=3).value = f"=CAPA!D{capa_r}"
        ws_4b.cell(row=row_4b, column=4).value = f"='QUADRO II'!D{row_q2}"
        ws_4b.cell(row=row_4b, column=5).value = f"='QUADRO II'!F{row_q2}"
        ws_4b.cell(row=row_4b, column=6).value = f"=('QUADRO II'!K{row_q2})+('QUADRO II'!R{row_q2})"
        ws_4b.cell(row=row_4b, column=7).value = f"=SUM(E{row_4b}+F{row_4b})"
        ws_4b.cell(row=row_4b, column=8).value = f"=('QUADRO II'!N{row_q2})"
        ws_4b.cell(row=row_4b, column=9).value = f"=('QUADRO II'!V{row_q2})"

    # Deletar linhas excedentes do IV B
    excesso_4b = total_ivb_slots - n
    if excesso_4b > 0:
        ws_4b.delete_rows(IVB_FIRST + n, excesso_4b)

    # ── 5c. QUADRO IV B.1 ─────────────────────────────────────────────────────
    ws_4b1 = wb["QUADRO IV B.1"]
    # Template: linha 15=UH1 (linha 14 é cabeçalho A B C D...), linhas 16+ com #REF!
    IVB1_FIRST = 15

    # Contar slots disponíveis
    ivb1_slots = 0
    for r in range(IVB1_FIRST + 1, 200):
        row_cells = list(ws_4b1.iter_rows(min_row=r, max_row=r))[0]
        has_border = any(c.border.left.border_style or c.border.right.border_style
                         or c.border.top.border_style or c.border.bottom.border_style
                         for c in row_cells)
        has_content = any(c.value is not None for c in row_cells)
        if not has_border and not has_content:
            break
        ivb1_slots += 1
    total_ivb1_slots = 1 + ivb1_slots

    for i in range(n):
        row_b1 = IVB1_FIRST + i
        row_q2 = Q2_TPL + i
        capa_r = CAPA_TPL + i
        row_4b = IVB_FIRST + i   # linha correspondente no IV B (após deleção, mesmo índice)
        ws_4b1.cell(row=row_b1, column=2).value  = f"=('QUADRO II'!B{row_q2})"
        ws_4b1.cell(row=row_b1, column=3).value  = f"=CAPA!D{capa_r}"
        ws_4b1.cell(row=row_b1, column=4).value  = f"='QUADRO IV B'!D{row_4b}"
        ws_4b1.cell(row=row_b1, column=5).value  = f"=SUM(C{row_b1}+D{row_b1})"
        ws_4b1.cell(row=row_b1, column=6).value  = f"=('QUADRO II'!K{row_q2})+('QUADRO II'!R{row_q2})"
        ws_4b1.cell(row=row_b1, column=7).value  = f"=SUM(E{row_b1}+F{row_b1})"
        ws_4b1.cell(row=row_b1, column=8).value  = f"=G{row_b1}"
        ws_4b1.cell(row=row_b1, column=10).value = f"=SUM(H{row_b1}+I{row_b1})"
        ws_4b1.cell(row=row_b1, column=11).value = f"=('QUADRO II'!N{row_q2})"
        ws_4b1.cell(row=row_b1, column=12).value = f"=('QUADRO II'!V{row_q2})"

    # Deletar linhas excedentes do IV B.1
    excesso_4b1 = total_ivb1_slots - n
    if excesso_4b1 > 0:
        ws_4b1.delete_rows(IVB1_FIRST + n, excesso_4b1)

    # ── 6. QUADRO V — limpar linhas de grupos extras com #REF! ──────────────
    # O template tem linhas L27=grupo1, L28..L32=grupos extras (para empreendimentos
    # com múltiplos tipos de UH). Para empreendimentos com 1 tipo, limpar L28:L32.
    ws_5 = wb["QUADRO V"]
    # L27 referencia QUADRO II!F17 (1ª UH) e CAPA!H9 — já corretos
    # L28:L32 têm #REF! do template — limpar valores (manter texto B/E/G intacto)
    for r_v in range(28, 33):
        for col_v in ["D", "F"]:
            cell_v = ws_5[f"{col_v}{r_v}"]
            if cell_v.value is not None and "#ref" in str(cell_v.value).lower():
                cell_v.value = None

    wb.save(out)
    print(f"  Quadros ABNT: {out}")
    return out


# ── Main ─────────────────────────────────────────────────────────────────────

def main():
    print("=" * 60)
    print("fetch_empreendimentos.py — Morais Engenharia  v3")
    print("=" * 60)

    # 1. Busca empreendimentos
    paginas_emp = buscar_paginas(NOTION_DB, HEADERS, "EMPREENDIMENTOS")

    # 2. Busca unidades
    paginas_uh = buscar_paginas(NOTION_DB_UH, HEADERS_UH, "UNIDADES DO EMPREENDIMENTO")
    unidades = []
    for p in paginas_uh:
        try:
            unidades.append(processar_unidade(p))
        except Exception as ex:
            print(f"  ERRO UH {p.get('id','?')}: {ex}")

    print(f"\nUnidades processadas: {len(unidades)}")

    # 3. Processa empreendimentos e gera documentos
    empreendimentos = []
    for p in paginas_emp:
        try:
            emp = processar_pagina(p)
            if not emp["nome"]:
                print(f"  IGNORADO: sem nome (id={p.get('id','?')})"); continue

            empreendimentos.append(emp)
            print(f"\n{'─'*50}")
            print(f"  {emp['nome']} — {emp['status_label']}  (RV{emp['revisao']})")

            gerar_fre(emp)
            gerar_carta_proposta(emp)
            gerar_memorial(emp, "HABITACAO")
            gerar_memorial(emp, "INFRAESTRUTURA")
            gerar_quadros_abnt(emp, unidades)

        except Exception as ex:
            import traceback
            print(f"  ERRO {p.get('id','?')}: {ex}")
            traceback.print_exc()

    empreendimentos.sort(key=lambda x: x["nome"].lower())

    saida = {
        "updated_at": datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ"),
        "empreendimentos": empreendimentos,
    }
    with open("data_emp.json", "w", encoding="utf-8") as f:
        json.dump(saida, f, ensure_ascii=False, indent=2)

    print(f"\n{'=' * 60}")
    print(f"data_emp.json: {len(empreendimentos)} empreendimento(s)")
    print("=" * 60)

if __name__ == "__main__":
    main()
